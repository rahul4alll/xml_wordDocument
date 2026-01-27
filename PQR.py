from lxml import etree
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
from docx.shared import Inches
from docx.shared import RGBColor
from html import unescape

LOGIC_RED = RGBColor(255, 0, 0)
INFO_BLUE = RGBColor(68, 114, 196)
GRAY_TEXT = RGBColor(128, 128, 128)

# =========================
# GLOBAL CONSTANTS
# =========================


EXPORT_START_LABEL = "te1"
EXPORT_END_LABEL = "b3"

EXPORT_ENABLED = False

# =========================
# ENTRY POINT (FILE BASED)
# =========================
def generate_word_from_xml_file(xml_path, output_path):
    parser = etree.XMLParser(recover=True)
    tree = etree.parse(xml_path, parser)
    root = tree.getroot()

    doc = Document()

    # =========================
    # HELPERS
    # =========================
    def local(tag):
        return etree.QName(tag).localname.lower()

    def safe(text):
        return text.strip() if text else ""

    # =========================
    # SURVEY NAME DETECTION (ALT FIX)
    # =========================
    def get_survey_name(root, default="Survey Specification Document"):
        # 1. Root <survey alt="...">
        if root.get("alt"):
            return root.get("alt").strip()

        # 2. <survey alt="..."> anywhere
        survey_node = root.xpath('.//*[local-name()="survey"]')
        if survey_node and survey_node[0].get("alt"):
            return survey_node[0].get("alt").strip()

        # 3. title / name / label attributes
        for attr in ("title", "name", "label"):
            if root.get(attr):
                return root.get(attr).strip()

        # 4. <title> node
        title_node = root.xpath('./*[local-name()="title"]')
        if title_node and title_node[0].text:
            return title_node[0].text.strip()

        return default


    INFO_BLUE = RGBColor(68, 114, 196)  # same blue as question header


    doc = Document()

    SURVEY_NAME = get_survey_name(root)
    heading = doc.add_heading(SURVEY_NAME, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER


    def red_text(text, style=None, bold_flag=False):
        p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
        r = p.add_run(text)
        r.font.color.rgb = RGBColor(255, 0, 0)
        r.bold = bold_flag
        return p

    def blue_text(text, style=None, bold_flag=False):
        p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
        r = p.add_run(text)
        r.font.color.rgb = RGBColor(0, 0, 255)
        r.bold = bold_flag
        return p

    # =========================
    # LEGEND / INSTRUCTIONS
    # =========================

    # Red bullet – Programming Logic
    p1 = doc.add_paragraph(style="List Bullet")
    r1 = p1.add_run("Red highlighted for Programming Logic")
    r1.bold = True
    r1.font.color.rgb = RGBColor(255, 0, 0)

    # Blue bullet – Instructions
    p2 = doc.add_paragraph(style="List Bullet")
    r2 = p2.add_run("Blue highlighted for Instructions")
    r2.bold = True
    r2.font.color.rgb = RGBColor(68, 114, 196)

    # =========================
    # QUESTION_TYPES
    # =========================
    QUESTION_TYPES = {"radio", "checkbox", "select", "text", "textarea", "number", "float"}

    # =========================
    # HELPERS
    # =========================

    LAST_ELEMENT_WAS_SUSPEND = False

    def extract_tooltip_from_text(text):

        if not text:
            return text, None, None

        text = unescape(text)

        pattern = re.compile(
            r'<span\s+class="tooltip">\s*(.*?)\s*'
            r'<span\s+class="tooltiptext">\s*(.*?)\s*</span>\s*</span>',
            flags=re.IGNORECASE | re.DOTALL
        )

        m = pattern.search(text)
        if not m:
            return text, None, None

        label = m.group(1).strip()
        definition = m.group(2).strip()

        clean_text = pattern.sub(label, text)
        clean_text = re.sub(r'\s+', ' ', clean_text).strip()

        return clean_text, label, definition

    def add_prefixed_rich_text(prefix, html_text):

        p = doc.add_paragraph()
        r = p.add_run(f"{prefix}: ")
        r.bold = True

        add_text_with_inline_html(p, html_text)

    def extract_tooltip_from_xml(title_elem):
        """
        Supports:
        1) Real XML tooltip spans
        2) Escaped HTML tooltip spans
        """

        # ---------- CASE 1: REAL XML TOOLTIP ----------
        tooltip = title_elem.xpath(
            './/*[local-name()="span" and @class="tooltip"]'
        )

        if tooltip:
            tooltip = tooltip[0]
            label = (tooltip.text or "").strip()

            tooltiptext = tooltip.xpath(
                './/*[local-name()="span" and @class="tooltiptext"]'
            )

            definition = (
                ''.join(tooltiptext[0].itertext()).strip()
                if tooltiptext else None
            )

            parts = []
            for node in title_elem.iter():
                if node is tooltip:
                    parts.append(label)
                elif node.getparent() is tooltip:
                    continue
                elif node.text:
                    parts.append(node.text)

            clean_text = ' '.join(' '.join(parts).split())
            return clean_text, label, definition

        # ---------- CASE 2: ESCAPED HTML TOOLTIP ----------
        raw_text = ''.join(title_elem.itertext()).strip()
        return extract_tooltip_from_text(raw_text)

    def resolve_definition_text(text, q):
        """
        Resolves ${res.X} → resource text
        """
        if not text:
            return None

        m = re.match(r'\$\{res\.(\w+)\}', text)
        if m:
            return get_resource_text(q, m.group(1))  # your existing function

        return text

    def get_resource_text(q, resource_name):
        # example implementation
        return q.get("resources", {}).get(resource_name, "")

    def hex_to_rgb(hex_color):
        hex_color = hex_color.lstrip("#")
        if len(hex_color) == 6:
            return RGBColor(
                int(hex_color[0:2], 16),
                int(hex_color[2:4], 16),
                int(hex_color[4:6], 16),
            )
        return None

    def add_text_with_inline_html(p, text):
        if not text:
            return

        text = text.replace("&nbsp;", " ")

        tokens = re.split(
            r'(</?(?:strong|b|i|em|u|span|li)[^>]*>|<br\s*/?>)',

            text,
            flags=re.IGNORECASE
        )

        bold = italic = underline = False
        color = None

        for token in tokens:
            t = token.lower().strip()

            if t in ("<b>", "<strong>"):
                bold = True
            elif t in ("</b>", "</strong>"):
                bold = False
            elif t in ("<i>", "<em>"):
                italic = True
            elif t in ("</i>", "</em>"):
                italic = False
            elif t == "<u>":
                underline = True
            elif t == "</u>":
                underline = False
            elif t.startswith("<br"):
                p.add_run("\n")
        # -------- Bullet handling --------
            elif t == "<li":
                # create a NEW bullet paragraph
                p = p._parent.add_paragraph(style="List Bullet")                
            elif t.startswith("<span"):
                m = re.search(r'color\s*:\s*(#[0-9a-fA-F]{6})', t)
                if m:
                    color = RGBColor.from_string(m.group(1).replace("#", ""))
            elif t == "</span>":
                color = None
            else:
                run = p.add_run(token)
                run.bold = bold
                run.italic = italic
                run.underline = underline
                if color:
                    run.font.color.rgb = color


    def render_list(elem):
        """Renders <ul><li>...</li></ul> as Word bullets"""
        for li in elem.xpath('./*[local-name()="li"]'):
            p = doc.add_paragraph(style="List Bullet")
            if li.text:
                add_text_with_inline_html(p, li.text)
        for c in li:
            if local(c.tag) == "br":
                p.add_run("\n")
            else:
                add_text_with_inline_html(p, c.text or "")
            if c.tail:
                add_text_with_inline_html(p, c.tail)

    def should_export(elem):
        """
        Controls export range between te1 and b3
        """
        global EXPORT_ENABLED

        label = elem.get("label")

        # Start marker (do NOT export te1)
        if label == EXPORT_START_LABEL:
            EXPORT_ENABLED = True
            return False

        # End marker (do NOT export b3)
        if label == EXPORT_END_LABEL:
            EXPORT_ENABLED = False
            return False

        return EXPORT_ENABLED


    def strip_hides_define_cond(q):
        """
        If question has strip="cond", suppress define-level condition display
        """
        return q is not None and q.get("strip", "").lower() == "cond"


    def get_display_cond(elem):
        """
        Returns display condition for both XML elements and define dicts
        """
        if isinstance(elem, dict):
            return elem.get("cond")
        return elem.get("cond")


    def add_layout_logic(q):
        """
        Displays keepWith / rightOf layout rules in Word export
        """
        keep_with = q.get("keepWith")
        right_of = q.get("rightOf")

        if not keep_with and not right_of:
            return

        p = doc.add_paragraph()

        if keep_with:
            r = p.add_run(f"Layout Logic: Keep with {keep_with} (Same Page)")
            r.bold = True
            r.font.color.rgb = LOGIC_RED

        if right_of:
            r = p.add_run(f"Layout Logic: Right of {right_of} (Same Page)")
            r.bold = True
            r.font.color.rgb = LOGIC_RED


    RES_VAR_PATTERN = re.compile(r"\$\{res\.([A-Za-z0-9_]+)\}")

    def resolve_res_value(text):
        """
        Replaces ${res.X} with actual <res label="X">value</res>
        """
        if not text:
            return text

        def replacer(match):
            var = match.group(1)
            return RES_VALUES.get(var, var)  # fallback to var name if missing

        return RES_VAR_PATTERN.sub(replacer, text)


    def resolve_uses_question_name(q):
        """
        Returns human-readable question name based on uses / atleast attributes
        """
        uses = q.get("uses", "")
        atleast = q.get("atleast")

        if not uses:
            return None

        uses = uses.lower()

        # Date
        if uses.startswith("fvdatepicker"):
            return "Date Question"

        # Card Rating
        if uses.startswith("cardrating"):
            return "Card Rating Question"

        # Slider variants
        if uses.startswith("sliderpoints"):
            return "Slider Rating Question"

        if uses.startswith("slidernumber"):
            return "Slider Question"

        if uses.startswith("sliderdecimal"):
            return "Slider Decimal Question"

        # Card Sort
        if uses.startswith("cardsort"):
            if atleast and atleast.isdigit() and int(atleast) > 1:
                return "Card Sort Multi Select Question"
            return "Card Sort Single Select Question"

        # Autosum
        if uses.startswith("autosum"):
            return "Autosum Question"

        # Rank Sort (multiple versions)
        if uses.startswith("ranksort"):
            return "Ranksort Question"

        # This That
        if uses.startswith("leftright"):
            return "This-That Question"

        if uses.startswith("imgmap"):
            return "Image highlighter Question"

        if uses.startswith("hottext"):
            return "Text highlighter Question"

        if uses.startswith("autosuggest"):
            return "Autosuggest Question"

        return None


    def local(tag):
        return etree.QName(tag).localname.lower()

    def show_optional_if_needed(q):
        qtype = local(q.tag)
        optional = q.get("optional")

        if qtype in {"number", "text", "textarea"}:
            # Explicit optional
            if optional == "1":
                show = True
            # Default optional for text / textarea
            elif qtype in {"text", "textarea"} and optional is None:
                show = True
            else:
                show = False

            if show:
                p = doc.add_paragraph()
                r = p.add_run("Optional Question")
                r.bold = True
                r.font.color.rgb = RGBColor(255, 0, 0)


    '''
    def add_separator_line():
        p = doc.add_paragraph()
        r = p.add_run("—" * 35)
        r.bold = True
    '''

    def get_row_text(row):
        if isinstance(row, dict):
            return safe(row.get("text"))
        return safe(row.text)


    def add_horizontal_line():
        p = doc.add_paragraph()

        p_pr = p._p.get_or_add_pPr()

        p_bdr = OxmlElement('w:pBdr')

        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')        # thickness
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')

        p_bdr.append(bottom)
        p_pr.append(p_bdr)


    from html import unescape

    def parse_groups(q):
        """
        Returns:
        groups: {group_label: clean_group_text}
        """
        groups = {}

        for g in q.xpath('./*[local-name()="group"]'):
            label = g.get("label")
            if not label:
                continue

            # Convert XML to HTML string
            text = etree.tostring(g, encoding="unicode", method="html")

            # Remove opening and closing <group> tags
            text = re.sub(r'</?group[^>]*>', '', text, flags=re.IGNORECASE)

            # Decode HTML entities (&lt;b&gt; → <b>)
            text = unescape(text)

            # Final safety cleanup (JUST IN CASE)
            text = text.replace('</group>', '').strip()

            groups[label] = text

        return groups


    ZERO_TOKEN = re.compile(r'(?<![\w.])0(?![\w.])')
    def is_hidden(elem):
        if elem is None:
            return False

        cond = elem.get("cond")
        if not cond:
            return False

        cond_clean = cond.strip().lower()

        # 1️⃣ Exact hidden
        if cond_clean == "0":
            return True

        # 2️⃣ Starts with "0 and ..."
        if re.match(r'^0\s*(and|&&)\b', cond_clean):
            return True

        # 3️⃣ Contains "... and 0 ..." (standalone zero)
        if re.search(r'\b(and|&&)\s*0\b', cond_clean):
            return True

        if re.search(r'(and|&&)\s*(?<![\w.])0(?![\w.])', cond_clean):
            return True

        return False


    def safe(text):
        return text.strip() if text else ""

    def get_any_cond(elem, cond_type):
        for k, v in elem.attrib.items():
            if k.lower() == cond_type.lower():
                return v
        return None

    def bold(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        return p

    def parse_exclude(exclude_value):

        if not exclude_value:
            return set()
        return {x.strip().lower() for x in exclude_value.split(",")}


    #from docx.shared import RGBColor

    def add_info_rich_text(elem):
        p = doc.add_paragraph()

        def walk(node):
            if node.text:
                #p.add_run(node.text)
                add_text_with_inline_html(p, node.text)

            for c in node:
                tag = local(c.tag)
                #run = p.add_run(c.text or "")
                add_text_with_inline_html(p, c.text or "")

                if tag in {"b", "strong"}:
                    run.bold = True
                elif tag == "i":
                    run.italic = True
                elif tag == "em":
                    run.italic = True
                elif tag == "u":
                    run.underline = True
                elif tag == "br":
                    p.add_run("\n")
                elif tag == "li":
                    p.style = "List Bullet"


                if c.tail:
                    p.add_run(c.tail)

        walk(elem)

    def get_attr(elem, *names):

        for k, v in elem.attrib.items():
            if "}" in k:
                lname = k.split("}", 1)[1]
            elif ":" in k:
                lname = k.split(":", 1)[1]
            else:
                lname = k

            if lname in names:
                return v
        return None


    def add_option_rich_text(elem, label_prefix, flags, cond_type=None, parent_q=None):

        p = doc.add_paragraph(style="List Continue")

        # Label
        r = p.add_run(f"{label_prefix}: ")
        r.bold = True

        # ===== XML ELEMENT =====
        if not isinstance(elem, dict):

            def walk(node):
                if node.text:
                    #p.add_run(node.text)
                    add_text_with_inline_html(p, node.text)


                for c in node:
                    tag = local(c.tag)
                    #run = p.add_run(c.text or "")
                    #add_text_with_inline_html(p, c.text or "")


                    if tag in {"b", "strong"}:
                        run = p.add_run(c.text or "")
                        run.bold = True
                    elif tag in {"i", "em"}:
                        run = p.add_run(c.text or "")
                        run.italic = True
                    elif tag == "u":
                        run = p.add_run(c.text or "")
                        run.underline = True
                    elif tag == "br":
                        p.add_run().add_break()
                    else:
                        add_text_with_inline_html(p, c.text or "")

                    if c.tail:
                        add_text_with_inline_html(p, c.tail)


            walk(elem)


            # Shuffle / Order logic (OPTION LEVEL)
            shuffle_logic = get_shuffle_logic(elem)
            for s in shuffle_logic:
                sr = p.add_run(f" ({s})")
                sr.bold = True
                sr.font.color.rgb = LOGIC_RED


        # ===== DEFINE / INSERT (DICT) =====
        else:
            add_text_with_inline_html(p, elem.get("text", ""))


        # Flags
        for f in flags:
            fr = p.add_run(f" - {f}")
            fr.bold = True
            fr.font.color.rgb = LOGIC_RED

        # Row / Col / Choice condition
        if cond_type:
            specific_cond = get_any_cond(elem, cond_type)
            if specific_cond:
                cr = p.add_run(f" ({cond_type}: {specific_cond})")
                cr.font.color.rgb = LOGIC_RED


        cond = get_display_cond(elem)

        # 🚫 Hide define-level cond when question has strip="cond"
        if cond:
            hide_define_cond = (
                parent_q
                and strip_hides_define_cond(parent_q)
                and isinstance(elem, dict)   # only define items
            )

            if not hide_define_cond:
                dr = p.add_run(f" (Display Condition: {cond})")
                dr.font.color.rgb = LOGIC_RED

        # Display condition
        #cond = elem.get("cond") if isinstance(elem, dict) else elem.get("cond")
        #if cond:
            #dr = p.add_run(f" (Display Condition: {cond})")
            #dr.font.color.rgb = LOGIC_RED

        # No Answer
        if is_noanswer(elem):
            nr = p.add_run(" (Exclusive)")
            nr.bold = True
            nr.font.color.rgb = LOGIC_RED


    # =========================
    # ANCHOR DETECTION
    # =========================
    ANCHOR_KEYWORDS = (
        "other", "none", "dk", "don't know",
        "dont know", "na", "n/a", "not applicable"
    )

    def is_anchor_text(text):
        t = text.lower().strip()
        return any(t.startswith(k) for k in ANCHOR_KEYWORDS)

    # =========================
    # NO ANSWER DETECTION (NEW)
    # =========================
    def is_noanswer(elem):
        if isinstance(elem, dict):
            return elem.get("noanswer") == "1"
        return local(elem.tag) == "noanswer" or elem.get("noanswer") in {"1", "true", "yes"}


    # =========================
    # GROUPS
    # =========================
    from collections import OrderedDict

    def group_rows(rows):
        """
        Returns:
        grouped_rows: OrderedDict {group_label: [rows]}
        ungrouped_rows: [rows]
        """
        grouped = OrderedDict()
        ungrouped = []

        for r in rows:
            grp = r.get("groups") if not isinstance(r, dict) else r.get("groups")
            if grp:
                for g in grp.split(","):
                    grouped.setdefault(g.strip(), []).append(r)
            else:
                ungrouped.append(r)

        return grouped, ungrouped

    def group_rows_by_group(rows):
        grouped = {}
        ungrouped = []

        for r in rows:
            grp = r.get("groups") if not isinstance(r, dict) else r.get("groups")
            if grp:
                for g in grp.split(","):
                    grouped.setdefault(g.strip(), []).append(r)
            else:
                ungrouped.append(r)

        return grouped, ungrouped

    def get_row_text(r):
        return r.get("text") if isinstance(r, dict) else safe(r.text)
    # =========================
    # DEFINES
    # =========================
    DEFINES = {}

    for d in root.xpath('.//*[local-name()="define"]'):
        label = d.get("label")
        if not label:
            continue

        items = []
        for r in d.xpath(
            './/*[local-name()="row" or local-name()="col" '
            'or local-name()="choice" or local-name()="value" '
            'or local-name()="noanswer"]'
        ):
            items.append({
                "tag": local(r.tag),
                "label": r.get("label", ""),
                "text": safe(r.text),
                "randomize": r.get("randomize"),
                "exclusive": r.get("exclusive"),
                "cond": r.get("cond"),
                "rowCond": r.get("rowCond"),
                "colCond": r.get("colCond"),
                "choiceCond": r.get("choiceCond"),
                "noanswer": "1" if local(r.tag) == "noanswer" else r.get("exclusive")
            })

        DEFINES[label] = items

    # =========================
    # RES VALUES (res resolver)
    # =========================
    RES_VALUES = {}

    for r in root.xpath('.//*[local-name()="res"]'):
        label = r.get("label")
        value = safe("".join(r.itertext()))
        if label:
            RES_VALUES[label] = value

    # =========================
    # RICH TEXT
    # =========================
    def add_rich_text(elem, prefix=None):
        p = doc.add_paragraph()
        if prefix:
            r = p.add_run(prefix)
            r.bold = True

        def walk(node):
            if node.text:
                #p.add_run(node.text)
                add_text_with_inline_html(p, node.text)
            for c in node:
                #run = p.add_run(c.text or "")
                add_text_with_inline_html(p, c.text or "")
                # ✅ REAL LIST HANDLING
                if tag == "ul":
                    render_list(c)
                    continue

                tag = local(c.tag)
                if tag in {"b", "strong"}:
                    run.bold = True
                elif tag == "i":
                    run.italic = True
                elif tag == "em":
                    run.italic = True
                elif tag == "u":
                    run.underline = True
                elif tag == "br":
                    p.add_run("\n")
                if c.tail:
                    p.add_run(c.tail)

        walk(elem)

    # =========================
    # TERM / EXEC
    # =========================
    def add_flow(elem):
        if not should_export(elem):
            return
        if is_hidden(elem):
            return    

        tag = local(elem.tag)
        if tag == "term":
            cond = elem.get("cond", "").strip()

            p = doc.add_paragraph()
            text = "🚫 Terminate Logic"
            if cond:
                text += f" : {cond}"

            r = p.add_run(text)
            r.bold = True
            r.font.color.rgb = RGBColor(255, 0, 0)

            #add_horizontal_line()
    '''
    def add_flow(elem):
        tag = local(elem.tag)
        if tag == "term":
            p = doc.add_paragraph()
            r = p.add_run("🚫 Terminate Logic")
            r.bold = True
            r.font.color.rgb = RGBColor(255, 0, 0)

            cond = elem.get("cond")
            if cond:
                #doc.add_paragraph(f"cond: {cond}", style="List Continue")
                #red_text(f"{cond}", style="List Continue")
                pr = doc.add_paragraph(style="List Continue")
                rr = pr.add_run(cond)
                rr.font.color.rgb = RGBColor(255, 0, 0)
            add_horizontal_line()


        elif tag == "exec":
            bold("⚙ Execution Logic")
            if elem.text:
                #doc.add_paragraph(safe(elem.text))
                red_text(safe(elem.text))
    '''

    INFO_BLUE = RGBColor(68, 114, 196)

    def add_info(elem):
        global LAST_ELEMENT_WAS_SUSPEND

        if is_hidden(elem):
            LAST_ELEMENT_WAS_SUSPEND = False
            return

        if not should_export(elem):
            return

        tag = local(elem.tag)

        # =========================
        # HTML INFO
        # =========================
        if tag == "html":
            LAST_ELEMENT_WAS_SUSPEND = False

            label = elem.get("label", "").strip()
            cond = elem.get("cond")

            if label:
                p = doc.add_paragraph()
                r = p.add_run(f"{label} (🅘 Information)")
                r.bold = True
                r.font.color.rgb = INFO_BLUE

            if cond:
                p = doc.add_paragraph()
                r = p.add_run(f"Display Condition: {cond}")
                r.bold = True
                r.font.color.rgb = LOGIC_RED

            add_info_rich_text(elem)

        # =========================
        # SUSPEND (PAGE BREAK LOGIC)
        # =========================
        elif tag == "suspend":

            # 🚫 If suspend is hidden → ignore completely
            if is_hidden(elem):
                LAST_ELEMENT_WAS_SUSPEND = False
                return

            LAST_ELEMENT_WAS_SUSPEND = True

            p = doc.add_paragraph()
            r = p.add_run("—" * 35)
            r.bold = True
            #r.font.color.rgb = LOGIC_RED

            cond = elem.get("cond")
            if cond:
                pr = doc.add_paragraph()
                rr = pr.add_run(f"Display Condition: {cond}")
                rr.bold = True
                rr.font.color.rgb = LOGIC_RED

    # =========================
    # INSERT
    # =========================
    #def resolve_insert(elem):
        #return DEFINES.get(elem.get("source"), [])

    def resolve_insert(elem):
        source = elem.get("source")
        items = DEFINES.get(source, [])

        exclude_set = parse_exclude(elem.get("exclude"))

        if not exclude_set:
            return items

        filtered = []
        for item in items:
            label = item.get("label", "").lower()
            tag = item.get("tag")

            # Match r / c / ch prefixes
            if label and any(
                label == ex or label.startswith(ex)
                for ex in exclude_set
            ):
                continue

            filtered.append(item)

        return filtered

    # =========================
    # OPTION FORMAT
    # =========================
    def format_option(elem, text):
        flags = []
        anchor = is_anchor_text(text)

        randomize = elem.get("randomize") if not isinstance(elem, dict) else elem.get("randomize")
        exclusive = elem.get("exclusive") if not isinstance(elem, dict) else elem.get("exclusive")
        open_flag = elem.get("open") if not isinstance(elem, dict) else elem.get("open")

        if open_flag == "1" or randomize == "0" or is_anchor_text(text):
            flags.append("anchor")

        # Exclusive
        if exclusive == "1":
            flags.append("exclusive")

        return text, flags, False


    def write_option(elem, text, flags, cond_type=None):
        p = doc.add_paragraph(style="List Continue")

        # Option text (normal)
        p.add_run(text)

        for f in flags:
            r = p.add_run(f" - {f}")
            if f in {"anchor", "exclusive"}:
                r.bold = True
                r.font.color.rgb = LOGIC_RED

        # Row / Col / Choice condition
        if cond_type:
            specific_cond = get_any_cond(elem, cond_type)
            if specific_cond:
                r = p.add_run(f" ({cond_type}: {specific_cond})")
                r.font.color.rgb = LOGIC_RED

        # Display condition
        cond = elem.get("cond") if not isinstance(elem, dict) else elem.get("cond")
        if cond:
            r = p.add_run(f" (Display Condition: {cond})")
            r.font.color.rgb = LOGIC_RED

        # No Answer
        if is_noanswer(elem):
            r = p.add_run(" (Exclusive)")
            r.bold = True
            r.font.color.rgb = LOGIC_RED

    def get_shuffle_logic(elem):
        """
        Reads shuffle / order attributes and returns readable logic text
        """
        logic = []

        shuffle = elem.get("shuffle")
        sortRows = elem.get("sortRows")
        rowShuffle = elem.get("rowShuffle")
        colShuffle = elem.get("colShuffle")


        if shuffle:
            shuffle = shuffle.lower()

            if shuffle == "rows":
                logic.append("Randomize Rows")
            elif shuffle == "cols":
                logic.append("Randomize Columns")
            elif shuffle == "choice":
                logic.append("Randomize Choices")
            elif shuffle == "rows,groups":
                logic.append("Randomize Rows and Groups")

        if rowShuffle:
            rowShuffle = rowShuffle.lower()

            if rowShuffle == "flip" and shuffle == "rows":
                logic.append("Flip Rows")
            elif rowShuffle == "rflip" and shuffle == "rows":
                logic.append("Reverse Flip Rows")
            elif rowShuffle == "rotate" and shuffle == "rows":
                logic.append("Rotate Options")
            elif rowShuffle == "rrotate" and shuffle == "rows":
                logic.append("Reverse Rotate Options")

        if colShuffle:
            colShuffle = colShuffle.lower()

            if colShuffle == "flip" and shuffle == "cols":
                logic.append("Flip Column")
            elif colShuffle == "rflip" and shuffle == "cols":
                logic.append("Reverse Flip Column")
            elif colShuffle == "rotate" and shuffle == "cols":
                logic.append("Rotate Options")
            elif colShuffle == "rrotate" and shuffle == "cols":
                logic.append("Reverse Rotate Options")

        if sortRows:
            sortRows = sortRows.lower()
        if sortRows == "asc":
            logic.append("Alphabatic Order")
        elif sortRows == "dsc":
            logic.append("Reverse Alphabatic Order")


        return logic
    # =========================
    # SORT OPTIONS (UPDATED)
    # =========================
    def sort_options(items):
        normal, anchor, noanswer = [], [], []

        for o in items:
            text = o["text"] if isinstance(o, dict) else safe(o.text)

            if is_noanswer(o):
                noanswer.append(o)
            elif is_anchor_text(text):
                anchor.append(o)
            else:
                normal.append(o)

        return normal + anchor + noanswer


    # ✅ Add page break ONLY if suspend was shown before this question
    if LAST_ELEMENT_WAS_SUSPEND:
        doc.add_page_break()
        LAST_ELEMENT_WAS_SUSPEND = False

    # =========================
    # QUESTION
    # =========================
    def render_question(q):
        if not should_export(q):
            return
        if is_hidden(q):
            return
        has_term = False

        title_elem = q.xpath('.//*[local-name()="title"]')
        if not title_elem:
            return
        title_elem = title_elem[0]

        # --- TOOLTIP LOGIC ---
        clean_text, label, definition = extract_tooltip_from_xml(title_elem)

        label = q.get("label", "NO_LABEL")
        qtype = local(q.tag).upper()

        uses_name = resolve_uses_question_name(q)
        if uses_name:
            doc.add_heading(f"{label} ({uses_name})", level=4)
        else:
            doc.add_heading(f"{label} ({qtype})", level=4)

        # Display condition
        if q.get("cond"):
            red_text(f"Display Condition: {q.get('cond')}")

        # Render question text
        add_rich_text(title_elem, "Question: ")

    
        # Add tooltip definition directly below question
        if label and definition:
            resolved_def = resolve_definition_text(definition, q)

            dp = doc.add_paragraph()
            dp.add_run("Definition: ").bold = True
            dp.add_run(resolved_def)

        

        # Resolve uses-based question name
    
    
        # Numeric (number / float) metadata
        add_numeric_metadata(q)

        # Optional flag
        show_optional_if_needed(q)

        # Layout logic (keepWith / rightOf)
        add_layout_logic(q)

        # Render comments / instructions
        comment = q.xpath('.//*[local-name()="comment"]')
        if comment:
            add_rich_text(comment[0], "Respondent Instruction: ")

        # Question-level row / col / choice conditions
        for cond_type, label in [
            ("rowCond", "Row Condition"),
            ("colCond", "Column Condition"),
            ("choiceCond", "Choice Condition"),
        ]:
            q_cond = get_any_cond(q, cond_type)
            if q_cond:
                p = doc.add_paragraph()
                r = p.add_run(f"{label}: {q_cond}")
                r.bold = True
                r.font.color.rgb = RGBColor(255, 0, 0)

        # Shuffle / Order logic
        shuffle_logic = get_shuffle_logic(q)
        for s in shuffle_logic:
            p = doc.add_paragraph()
            r = p.add_run(s)
            r.bold = True
            r.font.color.rgb = LOGIC_RED

        # Process rows, columns, choices
        rows, cols, choices = [], [], []

        for e in q.xpath(
            './/*[local-name()="row" or local-name()="col" '
            'or local-name()="choice" or local-name()="value" '
            'or local-name()="noanswer"]'
        ):
            if is_hidden(e):
                continue
            tag = local(e.tag)
            if tag == "row":
                rows.append(e)
            elif tag == "col":
                cols.append(e)
            else:
                choices.append(e)

        # Process <insert> elements
        for ins in q.xpath('./*[local-name()="insert"]'):
            resolved = resolve_insert(ins)
            for i in resolved:
                if i.get("cond") == "0":
                    continue
                tag = i.get("tag")
                if tag == "row":
                    rows.append(i)
                elif tag == "col":
                    cols.append(i)
                else:
                    choices.append(i)

        # Grouped rows
        groups = parse_groups(q)
        grouped_rows, ungrouped_rows = group_rows_by_group(rows)

        if groups:
            bold("Rows:")
            for g_label, g_title in groups.items():
                add_prefixed_rich_text("Group", g_title)
                g_rows = grouped_rows.get(g_label, [])
                for r_item in sort_options(g_rows):
                    label = r_item.get("label") if isinstance(r_item, dict) else r_item.get("label", "")
                    text = get_row_text(r_item)
                    _, flags, _ = format_option(r_item, text)
                    add_option_rich_text(r_item, label, flags, parent_q=q)

            # Ungrouped rows
            if ungrouped_rows:
                bold("Other Rows:")
                for r_item in sort_options(ungrouped_rows):
                    label = r_item.get("label") if isinstance(r_item, dict) else r_item.get("label", "")
                    text = get_row_text(r_item)
                    _, flags, _ = format_option(r_item, text)
                    add_option_rich_text(r_item, label, flags, parent_q=q)

        else:
            if rows:
                bold("Rows:")
                for r_item in sort_options(rows):
                    label = r_item.get("label") if isinstance(r_item, dict) else r_item.get("label", "")
                    text = get_row_text(r_item)
                    _, flags, _ = format_option(r_item, text)
                    add_option_rich_text(r_item, label, flags, parent_q=q)

        if cols:
            bold("Columns:")
            for c in sort_options(cols):
                label = c.get("label") if isinstance(c, dict) else c.get("label", "")
                _, flags, _ = format_option(c, c.get("text") if isinstance(c, dict) else safe(c.text))
                add_option_rich_text(c, label, flags)

        if choices:
            bold("Answer Options:")
            for o in sort_options(choices):
                label = o.get("label") if isinstance(o, dict) else o.get("label", "")
                _, flags, _ = format_option(o, o.get("text") if isinstance(o, dict) else safe(o.text))
                add_option_rich_text(o, label, flags)

        # Process flow and info elements
        for child in q:
            tag = local(child.tag)
            if tag == "term":
                has_term = True
                add_flow(child)
            elif tag == "exec":
                add_flow(child)
            elif tag in {"html", "suspend"}:
                add_info(child)



    def add_numeric_metadata(q):
        if local(q.tag) not in {"number", "float"}:
            return

        def red_bold(text):
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = True
            r.font.color.rgb = RGBColor(255, 0, 0)


        # Range from range="" OR verify="range(x,y)"
        range_value = None

        # 1️⃣ Direct range attribute
        if "range" in q.attrib and q.get("range"):
            range_value = q.get("range")

        # 2️⃣ Verify range(x,y)
        verify = q.get("verify")
        if not range_value and verify:
            m = re.search(r"range\s*\(\s*([^)]+)\s*\)", verify)
            if m:
                range_value = m.group(1)

        if range_value:
            red_bold(f"Range - ({range_value})")

        # Post Text
        post_text = get_attr(q, "postText")
        if post_text is not None:
            resolved_post = resolve_res_value(post_text)
            red_bold(f"Post Text: {resolved_post}")

        # Pre Text
        pre_text = get_attr(q, "preText")
        if pre_text is not None:
            resolved_pre = resolve_res_value(pre_text)
            red_bold(f"Pre Text: {resolved_pre}")


        # Optional
        if q.get("optional") == "1":
            red_bold("Optional Question")




    # =========================
    # LOOP ITERATION RESOLVER
    # =========================
    def get_loop_iterations(loop):
        """
        Returns list of tuples:
        (display_text, cond)
        Supports MULTIPLE loopvar per looprow
        """
        iterations = []

        looprows = loop.xpath('./*[local-name()="looprow"]')
        if looprows:
            for lr in looprows:
                cond = lr.get("cond")

                vars_text = []
                for lv in lr.xpath('./*[local-name()="loopvar"]'):
                    name = lv.get("name", "").strip()
                    value = safe(lv.text)
                    if name and value:
                        vars_text.append(f"{name} = {value}")
                    elif value:
                        vars_text.append(value)

                combined_text = " | ".join(vars_text) if vars_text else ""

                label = lr.get("label")
                display = f"{label}. {combined_text}" if label else combined_text

                iterations.append((display, cond))

            return iterations

        # Fallback (define-based loops)
        source = loop.get("source")
        if source in DEFINES:
            for item in DEFINES[source]:
                display = f"{item['label']}: {item['text']}"
                iterations.append((display, item.get("cond")))
            return iterations

        iterations.append(("⚠ Dynamic loop (resolved at runtime)", None))
        return iterations


        # Case 2: define-based loops
        source = loop.get("source")
        if source in DEFINES:
            for item in DEFINES[source]:
                display = f"{item['label']}: {item['text']}"
                iterations.append((display, item.get("cond")))
            return iterations

        # Fallback
        iterations.append(("⚠ Dynamic loop (resolved at runtime)", None))
        return iterations


    # =========================
    # LOOP
    # =========================
    def render_loop(loop):
        if not should_export(loop):
            return
        if is_hidden(loop):
            return

        label = loop.get("label", "LOOP")
        doc.add_heading(f"🔁 Loop: {label}", level=2)

        # Loop display condition
        loop_cond = loop.get("cond")
        if loop_cond:
            red_text(f"Loop Display Condition: {loop_cond}")

        # Loop title
        title = loop.xpath('./*[local-name()="title"]')
        if title:
            add_rich_text(title[0], "Loop Title: ")

        # =========================
        # LOOP ITERATIONS (FIRST)
        # =========================
        bold("Loop Iterations:")
        for i, (text, cond) in enumerate(get_loop_iterations(loop), 1):
            p = doc.add_paragraph(f"{i}. {text}", style="List Continue")
            if cond:
                r = p.add_run(f" (Condition: {cond})")
                r.bold = True
                r.font.color.rgb = LOGIC_RED

        # =========================
        # LOOP CONTENT (SECOND)
        # =========================
        bold("Loop Content:")

        for child in loop:
            tag = local(child.tag)

            if tag == "block":
                render_block(child, in_loop=True)

            elif tag == "loop":
                render_loop(child)   # 🔁 RECURSION (NESTED LOOP)

            elif tag in QUESTION_TYPES:
                render_question(child)

            elif tag in {"term", "exec"}:
                add_flow(child)

            elif tag in {"html", "suspend"}:
                add_info(child)

        bold(f"🔚 END LOOP: {label}")


    '''
        bold("Loop Questions:")
        for q in loop.xpath(
            './/*[local-name()="radio" or local-name()="checkbox" '
            'or local-name()="select" or local-name()="text" '
            'or local-name()="textarea" or local-name()="number" or local-name()="float"]'
        ):
            render_question(q)
    '''
    # =========================
    # BLOCK
    # =========================
    def render_block(b, in_loop=False):
        if not should_export(b):
            return

        if is_hidden(b):
            return

        label = b.get("label", "BLOCK")
        if in_loop:
            bold(f"📦 START LOOP BLOCK: {label}")
        else:
            bold(f"📦 START BLOCK: {label}")


        block_cond = b.get("cond")
        if block_cond:
            if in_loop:
                red_text(f"Loop iteration logic: {block_cond}")
            else:
                red_text(f"Block Display Condition: {block_cond}")

        doc.add_heading(f"Block: {label}", level=3)

        for child in b:
            tag = local(child.tag)
            if tag == "block":
                render_block(child)
            if tag == "loop":
                render_loop(child)
            elif tag in QUESTION_TYPES:
                render_question(child)
            elif tag in {"term", "exec"}:
                add_flow(child)
            elif tag in {"html", "suspend"}:
                add_info(child)


        if in_loop:
            bold(f"📦 END LOOP BLOCK: {label}")
        else:
            bold(f"📦 END BLOCK: {label}")

    for elem in root:
        tag = local(elem.tag)
        if tag == "block":
            render_block(elem)
        elif tag == "loop":
            render_loop(elem)
        elif tag in QUESTION_TYPES:
            render_question(elem)
        elif tag in {"term", "exec"}:
            add_flow(elem)
        elif tag in {"html", "suspend"}:
            add_info(elem)


    doc.save(output_path)

